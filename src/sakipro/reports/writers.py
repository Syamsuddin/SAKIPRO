import os
import pandas as pd
from datetime import datetime
import docx

def export_to_markdown(content: str, command_name: str, output_dir: str = "output"):
    """Export text content to a markdown file."""
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{timestamp}_{command_name}.md"
    filepath = os.path.join(output_dir, filename)
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
        
    return filepath

def export_to_excel(data: list[dict], command_name: str, output_dir: str = "output"):
    """Export a list of dictionaries to an Excel file."""
    if not data:
        return None
        
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{timestamp}_{command_name}.xlsx"
    filepath = os.path.join(output_dir, filename)
    
    df = pd.DataFrame(data)
    df.to_excel(filepath, index=False)
    
    return filepath

def export_to_docx(content: str, command_name: str, output_dir: str = "output"):
    """Export text content to a Word Document."""
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{timestamp}_{command_name}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc = docx.Document()
    doc.add_heading(f"Laporan AI: {command_name}", 0)
    
    # Simple markdown to text conversion
    for line in content.split('\n'):
        if line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=2)
        else:
            doc.add_paragraph(line)
            
    doc.save(filepath)
    return filepath

